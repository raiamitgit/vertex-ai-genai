"""
ADK Agent: Document Processing and Image Generation

Shows how to build an ADK agent that handles document workflows, file interception, and image generation.

Key Functionalities:
1.  **JSON File Interception**: Use `before_model_callback` to intercept JSON files and inject content as text into the prompt.
2.  **Artifact Generation (Word Docs)**: Use a tool to generate `.docx` files and save as artifacts for Gemini Enterprise.
3.  **Image Generation**: Use `gemini-3.1-flash-image-preview` with `response_modalities=["IMAGE"]` to generate charts and save as artifacts.
4.  **Guardrails**: Enforce a file-centric focus where the agent refuses to answer questions unless a file is provided.
5.  **GCS Artifact Management**: Use Google Cloud Storage for artifacts to work with the Enterprise UI.
"""
import io
import logging
import mimetypes
import os
from pathlib import Path
import zipfile

from dotenv import load_dotenv

from google.adk.agents import Agent
from google.adk.agents.callback_context import CallbackContext
from google.adk.artifacts import FileArtifactService, GcsArtifactService
from google.adk.models import Gemini
from google.adk.models.llm_request import LlmRequest
from google.adk.runners import Runner
from google.adk.sessions import InMemorySessionService, VertexAiSessionService
from google.adk.tools import FunctionTool, ToolContext
from google.adk.tools.load_artifacts_tool import load_artifacts_tool
from google.genai import types

# Load environment variables from .env file in the same directory
load_dotenv(dotenv_path=Path(__file__).parent / ".env")

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Custom class to force global location for Gemini 3 models
class GlobalGemini(Gemini):
    def __init__(self, model: str):
        # Set the environment variable to force the ADK SDK to use the global location for the client.
        os.environ["GOOGLE_CLOUD_LOCATION"] = "global"
        super().__init__(model=model)

# --- Single Agent for Document Processing ---
SYSTEM_INSTRUCTION = """You are a Document Assistant. Your job is to assist users by answering questions derived from the files they upload.

### Rule: Respond with a Document, Image, or Video
For every response you generate (whether it's a summary, an answer to a question, or an analysis), you must not provide the answer in the text response directly. Instead, you must:

1. **For Text Answers & Reports**: If the user asks a question or requests a summary:
    - Generate the full answer content.
    - Invoke the `create_response_document` tool with the content and a filename (e.g., 'response.docx').
    - Your final text response to the user must only be a confirmation message like: "I have generated the response in a Word document for you."

2. **For Images & Infographics**: If the user asks to generate an image, chart, or infographic to explain something:
    - Generate a prompt describing the visual you want.
    - Invoke the `generate_infographic` tool with the prompt and a filename (e.g., 'chart.png').
    - Your final text response to the user must only be a confirmation message like: "I have generated the visuals for you."

3. **For Videos**: If the user asks to generate a video from text:
    - Generate a prompt describing the video you want.
    - Invoke the `generate_video` tool with the prompt.
    - Your final text response to the user should confirm that the video was generated and provide the GCS URI returned by the tool.

### Supported File Types:
1. **PDF Documents**: When a user references a PDF, use the `load_artifacts_tool` to extract the text.
2. **CSV Data Files**: When a user references a CSV, use the `load_artifacts_tool` to read the data.
3. **Images (PNG, JPG, JPEG)**: You can process image files natively.
4. **Videos (MP4)**: You can process video files natively by referencing their location.
5. **JSON Files**: You can process JSON files. The system will extract the content and provide it to you as text.
6. **Other Files**: If the user uploads any other file type, inform them that you cannot process that file type.

### Guardrails:
* **File-Centric Focus**: Answer questions based on the data provided in the uploaded files.
* **Enforce File Uploads**: If the user asks a question but has not uploaded any file, instruct them to upload one first.
* **No General Knowledge Queries**: Refuse to answer general questions unrelated to an uploaded file.
* **Always Use Tools**: Do not assume you know the content of a file without reading it.
* **No Code Execution**: Do not attempt to write or execute Python code to process files. Rely on the provided tools.

Keep interactions helpful within these boundaries.
"""

async def intercept_json_callback(callback_context: CallbackContext, llm_request: LlmRequest) -> None:
    """Callback executed before the model runs to intercept JSON files.
    
    RATIONALE:
    Intercept JSON files to ensure content is fully accessible as text in the prompt
    rather than relying on the model to extract it via tool calls.
    Guarantees content is read before generation.
    """
    logger.debug("Entering before_model_callback")
    if llm_request.contents:
        for i, content in enumerate(llm_request.contents):
            if content.role == "user":
                new_parts = []
                content_modified = False
                for part in content.parts:
                    # Extract file_data and inline_data, handling both objects and dictionaries
                    # (The UI may send these as dictionaries instead of Pydantic objects in some versions)
                    file_data = getattr(part, 'file_data', None)
                    inline_data = getattr(part, 'inline_data', None)
                        
                    mime_type = None
                    
                    if file_data:
                        mime_type = getattr(file_data, 'mime_type', None)
                        if not mime_type and isinstance(file_data, dict):
                            mime_type = file_data.get('mime_type')
                    elif inline_data:
                        mime_type = getattr(inline_data, 'mime_type', None)
                        if not mime_type and isinstance(inline_data, dict):
                            mime_type = inline_data.get('mime_type')
                            
                    # Intercept JSON files and inject content as text
                    if mime_type == "application/json":
                        logger.info("Found JSON file in request! Intercepting.")
                        content_modified = True
                        if file_data:
                            # File data indicates the file is in the artifact store (e.g. GCS)
                            filename = getattr(file_data, 'display_name', None) or os.path.basename(getattr(file_data, 'file_uri', ''))
                            if not filename and isinstance(file_data, dict):
                                filename = file_data.get('display_name') or os.path.basename(file_data.get('file_uri', ''))
                            logger.info(f"Attempting to load artifact: {filename}")
                            
                            # We use callback_context to load the artifact bytes
                            try:
                                artifact_part = await callback_context.load_artifact(filename)
                                json_bytes = artifact_part.inline_data.data
                                new_parts.append(types.Part.from_text(text=f"[System: Content of {filename}:]\n{json_bytes.decode('utf-8')}"))
                            except Exception as e:
                                logger.error(f"Failed to load artifact {filename}: {e}")
                                # Keep the original part so the model can try to use tools or report error
                                new_parts.append(part)
                                new_parts.append(types.Part.from_text(text=f"[System: Warning: Could not preload {filename}. Error: {e}]"))
                        else:
                            # Inline data means the file content came directly in the request
                            json_bytes = getattr(inline_data, 'data', None)
                            if not json_bytes and isinstance(inline_data, dict):
                                json_bytes = inline_data.get('data')
                            new_parts.append(types.Part.from_text(text=f"[System: Content of JSON file:]\n{json_bytes.decode('utf-8')}"))
                        continue
                        
                    # For all other parts, keep them as is
                    new_parts.append(part)
                
                if content_modified:
                    # Replace the entire content object to avoid Pydantic mutation errors
                    llm_request.contents[i] = types.Content(role="user", parts=new_parts)
                    logger.info("Handled special files in request.")
    return None

async def create_response_document(content: str, filename: str, tool_context: ToolContext) -> dict:
    """Creates a Word document with provided content and saves as an artifact.
    
    RATIONALE:
    Allows the agent to deliver findings as a file rather than chat text.
    Artifacts saved via `tool_context.save_artifact` are rendered in the UI.
    
    Args:
        content: Text content to put in the document.
        filename: Desired filename for the Word document (e.g., 'response.docx').
        tool_context: ADK ToolContext provided automatically.
    """
    from docx import Document
    import io
    
    doc = Document()
    doc.add_heading('Response', 0)
    
    # Simple markdown parsing for headings to give the document some structure
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.strip():
            doc.add_paragraph(line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create a Part object with the correct MIME type for Word documents
    part = types.Part(inline_data=types.Blob(
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buffer.read()
    ))
    
    # Save the artifact using ToolContext. This handles storage in GCS or local disk
    try:
        version = await tool_context.save_artifact(filename, part)
        
        return {
            "status": "success",
            "message": f"Word document saved as artifact: {filename}",
            "version": version
        }
    except Exception as e:
        logger.error(f"Failed to save artifact {filename}: {e}")
        return {
            "status": "error",
            "message": f"Failed to save artifact {filename}. Error: {e}"
        }

async def generate_infographic(prompt: str, filename: str, tool_context: ToolContext) -> dict:
    """Generates an image or infographic based on the prompt using Gemini 3.1 Flash Image Preview model.
    
    RATIONALE:
    Shows how to do text-to-image generation in an agent.
    Uses `gemini-3.1-flash-image-preview` by requesting the "IMAGE" response modality.
    
    Args:
        prompt: Description of the image to generate.
        filename: Desired filename for the saved image (e.g., 'infographic.png').
        tool_context: ADK ToolContext provided automatically.
    """
    from google.genai import Client
    from google.genai import types
    import os
    
    model_name = "gemini-3.1-flash-image-preview" 
    
    # Initialize client (using global region as requested)
    # Project is automatically picked up from GOOGLE_CLOUD_PROJECT environment variable
    client = Client(vertexai=True, location="global")
    
    try:
        logger.info(f"Attempting to generate image with model {model_name} and prompt: {prompt}")
        
        # To generate images with this preview model, we set response_modalities to ["IMAGE"]
        generate_content_config = types.GenerateContentConfig(
            response_modalities = ["IMAGE"],
        )
        
        response = client.models.generate_content(
            model=model_name,
            contents=[
                types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=prompt)]
                )
            ],
            config=generate_content_config,
        )
        
        # Extract the image bytes from the response candidates
        part = response.candidates[0].content.parts[0]
        
        if part.inline_data:
            image_bytes = part.inline_data.data
        else:
            raise ValueError("No inline data found in the response part.")
            
        save_part = types.Part(inline_data=types.Blob(
            mime_type="image/png",
            data=image_bytes
        ))
        
        # Save the generated image as an artifact
        version = await tool_context.save_artifact(filename, save_part)
        
        return {
            "status": "success",
            "message": f"Image saved as artifact: {filename}",
            "version": version
        }
    except Exception as e:
        logger.error(f"Failed to generate image: {e}")
        return {"status": "error", "message": str(e)}

async def generate_video(prompt: str, tool_context: ToolContext) -> dict:
    """Generates a video based on the prompt using Veo 3.1 Lite model.
    
    Args:
        prompt: Description of the video to generate.
        tool_context: ADK ToolContext provided automatically.
    """
    from google.genai import Client
    from google.genai.types import GenerateVideosConfig, GenerateVideosSource, Part, Blob
    import time
    import os
    
    model_name = "veo-3.1-fast-generate-001"
    bucket_name = os.getenv("GCS_ARTIFACT_BUCKET")
    
    if not bucket_name:
        return {"status": "error", "message": "GCS_ARTIFACT_BUCKET environment variable not set."}
        
    project_id = os.getenv("GOOGLE_CLOUD_PROJECT")
    client = Client(vertexai=True, project=project_id, location="us-central1")
    
    try:
        logger.info(f"Attempting to generate video with model {model_name} and prompt: {prompt}")
        
        output_gcs_uri = f"gs://{bucket_name}/veo_output/"
        
        source = GenerateVideosSource(prompt=prompt)
        
        operation = client.models.generate_videos(
            model=model_name,
            source=source,
            config=GenerateVideosConfig(
                aspect_ratio="16:9",
                number_of_videos=1,
                duration_seconds=8,
                resolution="720p",
            ),
        )
        
        logger.info(f"Video generation started. Operation: {operation.name}")
        
        # Polling for completion
        while not operation.done:
            logger.debug("Waiting for video generation operation...")
            time.sleep(15)
            operation = client.operations.get(operation)
            
        logger.info("Video generation completed.")
        logger.debug(f"Operation: {operation}")
        
        if operation.response:
            result = getattr(operation, 'result', None) or operation.response
            generated_videos = getattr(result, 'generated_videos', None)
            
            if generated_videos:
                video = generated_videos[0].video
                video_uri = getattr(video, 'uri', None)
                
                if video_uri:
                    return {
                        "status": "success",
                        "message": "Video generated successfully.",
                        "video_uri": video_uri
                    }
                
                # Check if it returned inline bytes or content
                video_bytes = getattr(video, 'video_bytes', None) or getattr(video, 'bytes', None) or getattr(video, 'inline_data', None)
                
                if video_bytes:
                    # If it's base64 string, we might need to decode it. 
                    if isinstance(video_bytes, str):
                        import base64
                        video_bytes = base64.b64decode(video_bytes.encode('utf-8'))
                    
                    filename = f"generated_video_{int(time.time())}.mp4"
                    part = Part(inline_data=Blob(
                        mime_type="video/mp4",
                        data=video_bytes
                    ))
                    
                    version = await tool_context.save_artifact(filename, part)
                    
                    return {
                        "status": "success",
                        "message": f"Video generated successfully and saved as artifact {filename}.",
                        "artifact_name": filename,
                        "version": version
                    }
                
                return {"status": "error", "message": f"Video generated but no URI or bytes found. Video object: {video}"}
            else:
                return {"status": "error", "message": f"No generated videos in result. Result was: {result}"}
        else:
            return {"status": "error", "message": f"Operation failed or no response. Operation: {operation}"}
            
    except Exception as e:
        logger.error(f"Failed to generate video: {e}")
        return {"status": "error", "message": str(e)}

# Agent definition handling conversation and generating artifacts.
orchestrator = Agent(
    name="minimal_summarizer",
    description="Processes documents including PDFs and CSVs, and generates artifacts.",
    model=GlobalGemini("gemini-3-flash-preview"),
    generate_content_config=types.GenerateContentConfig(
        thinking_config=types.ThinkingConfig(
            include_thoughts=True
        )
    ),
    instruction=SYSTEM_INSTRUCTION,
    tools=[load_artifacts_tool, create_response_document, generate_infographic, generate_video],
    before_model_callback=intercept_json_callback,
)

# Environment & Infrastructure Setup
running_in_cloud = os.getenv("RUNNING_IN_CLOUD", "false").lower() == "true"
project_id = os.getenv("GOOGLE_CLOUD_PROJECT")
location = os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1")
bucket_name = os.getenv("GCS_ARTIFACT_BUCKET")

if running_in_cloud:
    # In the cloud, use Vertex AI for sessions.
    print(f"Using VertexAiSessionService (Project: {project_id})")
    session_service = VertexAiSessionService(project=project_id, location=location)
else:
    # Locally, use in-memory sessions.
    print("Using InMemorySessionService (Local)")
    session_service = InMemorySessionService()

# Always use GCS for artifacts to ensure files are accessible in UI.
print(f"Using GcsArtifactService with bucket: {bucket_name}")
artifact_service = GcsArtifactService(bucket_name=bucket_name)

# Runner ties everything together.
runner = Runner(
    app_name="minimal_summarizer_agent",
    agent=orchestrator,
    session_service=session_service,
    artifact_service=artifact_service,
)

root_agent = runner.agent
