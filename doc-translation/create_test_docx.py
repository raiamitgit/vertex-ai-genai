from docx import Document

doc = Document()
doc.add_heading('Hello, World!', 0)
doc.add_paragraph('This is a simple test document for translation.')
doc.save('simple_test.docx')
print("Created simple_test.docx")
